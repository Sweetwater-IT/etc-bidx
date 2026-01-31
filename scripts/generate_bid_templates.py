from docx import Document

def main():
  # Estimator - Detailed
  doc = Document()
  doc.add_heading('ETC Sign Order Worksheet', 0)

  # Header table example
  table = doc.add_table(rows=1, cols=3)
  table.style = 'Table Grid'
  hdr_cells = table.rows[0].cells
  hdr_cells[0].text = 'Customer\n[Name]'
  hdr_cells[1].text = 'Job Number\n[Number]'
  hdr_cells[2].text = 'Contract #\n[#]'

  doc.add_heading('SIGN LIST', level=1)
  sign_table = doc.add_table(rows=1, cols=10)
  sign_table.style = 'Table Grid'
  headers = ['Designation', 'Description', 'Qty', 'Width', 'Height', 'Sheeting', 'Substrate', 'Stiffener', 'Unit Price', 'Total Price']
  for i, header in enumerate(headers):
    sign_table.rows[0].cells[i].text = header

  # Sample rows
  sample_rows = [
    ['DS1', 'High Intensity', '5', '4\'', '8\'', 'HI', 'Alum', 'Yes', '$100.00', '$500.00'],
    ['DS2', 'DG', '3', '3\'', '6\'', 'DG', 'Poly', 'No', '$80.00', '$240.00'],
    ['DS3', 'Special', '10', '5\'', '10\'', 'Special', 'Alum', 'Yes', '$150.00', '$1,500.00']
  ]
  for row in sample_rows:
    row_cells = sign_table.add_row().cells
    for i, text in enumerate(row):
      row_cells[i].text = text

  doc.save('public/documents/templates/bid-summary-estimator.docx')

  # PM Summary
  pm_doc = Document()
  pm_doc.add_heading('Bid Summary for Project Managers', 0)

  # Revenue table
  rev_table = pm_doc.add_table(rows=1, cols=3)
  rev_table.style = 'Table Grid'
  rev_table.rows[0].cells[0].text = 'Bid Item'
  rev_table.rows[0].cells[1].text = 'Total Revenue'
  rev_table.rows[0].cells[2].text = 'Percentage'

  rev_data = [
    ('MPT Mobilization', '$35,000.00', '35.00%'),
    ('MPT', '$65,000.00', '65.00%'),
    ('Rental', '$10,000.00', '10.00%'),
    ('Flagging', '$5,000.00', '5.00%'),
    ('Perm. Signs', '$3,000.00', '3.00%'),
    ('Sale', '$2,000.00', '2.00%'),
    ('Total', '$120,000.00', '100.00%')
  ]
  for data in rev_data:
    row_cells = rev_table.add_row().cells
    row_cells[0].text = data[0]
    row_cells[1].text = data[1]
    row_cells[2].text = data[2]

  pm_doc.add_heading('Discount Summary', level=1)
  disc_table = pm_doc.add_table(rows=1, cols=2)
  disc_table.style = 'Table Grid'
  disc_table.rows[0].cells[0].text = 'Item'
  disc_table.rows[0].cells[1].text = 'Rate'

  disc_data = [
    ('MPT', '20.00%'),
    ('SIGNS', '15.00%'),
    ('Total', '18.00%')
  ]
  for data in disc_data:
    row_cells = disc_table.add_row().cells
    row_cells[0].text = data[0]
    row_cells[1].text = data[1]

  pm_doc.save('public/documents/templates/bid-summary-pm.docx')

  print('Templates generated successfully!')

if __name__ == '__main__':
  main()
